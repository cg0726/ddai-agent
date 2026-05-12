from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT

from modules.config import BASE_DIR, EXPORT_DIR
from modules.project import get_sections, get_files, get_project


HEADER_TEXT = "中国银行广安分行 · 公司金融部"


def _set_font(run, font_name_cn: str, font_name_en: str, size: Pt, bold: bool = False):
    run.font.size = size
    run.font.bold = bold
    run.font.name = font_name_en
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = r.makeelement(qn("w:rPr"), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name_cn)
    rFonts.set(qn("w:ascii"), font_name_en)
    rFonts.set(qn("w:hAnsi"), font_name_en)


def _add_paragraph(doc, text: str, font_cn: str, font_en: str, size: Pt,
                   bold: bool = False, align: int = None, space_before: int = 0,
                   space_after: int = 0, line_spacing: float = 1.5):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    run = p.add_run(text)
    _set_font(run, font_cn, font_en, size, bold)
    return p


def _setup_page(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)


def _add_header(doc):
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run(HEADER_TEXT)
        _set_font(run, "仿宋", "FangSong", Pt(9), bold=False)


def _add_cover(doc, company_name: str, file_summary: str):
    for _ in range(6):
        _add_paragraph(doc, "", "仿宋", "FangSong", Pt(12))

    _add_paragraph(doc, "尽责调查报告", "黑体", "SimHei", Pt(28), bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    _add_paragraph(doc, "", "仿宋", "FangSong", Pt(12))

    today = datetime.now().strftime("%Y年%m月%d日")
    _add_paragraph(doc, f"项目名称：{company_name}", "仿宋", "FangSong", Pt(16),
                   align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    _add_paragraph(doc, f"生成日期：{today}", "仿宋", "FangSong", Pt(16),
                   align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)

    for _ in range(3):
        _add_paragraph(doc, "", "仿宋", "FangSong", Pt(12))

    _add_paragraph(doc, "上传文件清单", "黑体", "SimHei", Pt(16), bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    _add_paragraph(doc, "", "仿宋", "FangSong", Pt(8))
    _add_paragraph(doc, file_summary, "仿宋", "FangSong", Pt(12),
                   align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)

    doc.add_page_break()


def _add_toc(doc):
    _add_paragraph(doc, "目  录", "黑体", "SimHei", Pt(22), bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5,
                   space_before=200, space_after=200)
    for _ in range(3):
        _add_paragraph(doc, "", "仿宋", "FangSong", Pt(12))
    toc_text = "（请在 Word 中右键此处 → 更新域 以生成目录）\n" \
               "或按 Ctrl+A 后按 F9 刷新目录。"
    _add_paragraph(doc, toc_text, "仿宋", "FangSong", Pt(12),
                   align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fld_char1 = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    run._element.append(fld_char1)
    run2 = p.add_run()
    instr = run2._element.makeelement(qn("w:instrText"), {})
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    run2._element.append(instr)
    run3 = p.add_run()
    fld_char2 = run3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "separate"})
    run3._element.append(fld_char2)
    run4 = p.add_run("（目录，请在生成后更新域）")
    run5 = p.add_run()
    fld_char3 = run5._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run5._element.append(fld_char3)

    doc.add_page_break()


def export_to_word(project_id: int) -> Optional[str]:
    project = get_project(project_id)
    if not project:
        return None

    company_name = project["name"]
    sections = get_sections(project_id)
    confirmed_sections = [s for s in sections if s.get("confirmed")]
    files = get_files(project_id)

    if not confirmed_sections:
        return None

    file_summary_lines = []
    for f in files:
        file_summary_lines.append(f"  - {f['filename']} ({f['category']})")
    file_summary = "\n".join(file_summary_lines) if file_summary_lines else "无"

    doc = Document()
    _setup_page(doc)
    _add_header(doc)
    _add_cover(doc, company_name, file_summary)
    _add_toc(doc)

    for i, sec in enumerate(confirmed_sections):
        title = sec.get("title", f"第{i+1}章")
        content = sec.get("content", "")

        is_h1 = not title.startswith(("一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、", "九、", "十、"))
        if is_h1:
            _add_paragraph(doc, title, "黑体", "SimHei", Pt(16), bold=True,
                           align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5,
                           space_before=12, space_after=6)
        else:
            _add_paragraph(doc, title, "黑体", "SimHei", Pt(15), bold=True,
                           align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5,
                           space_before=8, space_after=4)

        if content:
            cleaned = content.strip()
            cleaned = cleaned.replace("\r\n", "\n").replace("\r", "\n")
            paragraphs = cleaned.split("\n")
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue
                is_bullet = para_text.startswith("- ") or para_text.startswith("• ")
                if is_bullet:
                    _add_paragraph(doc, para_text, "仿宋", "FangSong", Pt(12),
                                   align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5,
                                   space_before=2, space_after=2)
                else:
                    _add_paragraph(doc, para_text, "仿宋", "FangSong", Pt(12),
                                   align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5,
                                   space_before=0, space_after=0)

    today_str = datetime.now().strftime("%Y%m%d")
    safe_name = "".join(c if c.isalnum() or c in "_- " else "_" for c in company_name)
    filename = f"{safe_name}_尽责调查报告_{today_str}.docx"
    export_path = EXPORT_DIR / filename
    export_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(export_path))

    return str(export_path)
